from ast import keyword
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
import time
import docx
from docx.shared import Pt
from docx.oxml.ns import qn
from openpyxl import load_workbook
##待處理
####標題搜索連結加入
####屬性excel(勾選單複數在excel > 查題目自動兩種都查)並且在doc有目錄
####超連結不知怎麼調標楷體
####aminoisobutyric acid有BUG
####如果沒有非主題，產生更多題目

#---可調控變數---#
data="頸部考訓-解剖"
# data="1121-組織1"
PATH = 'C:/Users/tony/Desktop/Qbank/chromedriver.exe' #可變變數地區路徑
cooldown=1.0  #網速會影響


#---程式變數一覽---#
Qrepeatcheck=list() #檢測是否已經出現過這個題目
keywords=list() #存單字庫
checknum=0 #目前到第幾個單字
repeatnum=0 #某單字有N個題目重複
allquesnum=0 #某單字總共有N個題目
allnum=0 #所有單字總數
longquesnum=0 #某單字有N個長文題目
invalidtagnum=0 #某單字有N個不合法標籤題目
quetag=str() #某單字的標籤名稱
valid_tags={'解剖','組織','生理','病理','醫學','生物','藥劑','藥物','治療','臨床','構造','醫務','醫師','視力','聽力','麻醉','復健','醫生','藥理','診斷','護理'}
# baseurl="https://yamol.tw/tfulltext-"+keywords+".htm"

option = webdriver.ChromeOptions()
option.add_experimental_option("excludeSwitches", ["enable-automation"])
option.add_experimental_option('excludeSwitches', ['enable-logging'])
option.add_experimental_option('useAutomationExtension', False)
option.add_experimental_option("detach", True)
option.add_experimental_option("prefs", {"profile.password_manager_enabled": False, "credentials_enable_service": False})
driver = webdriver.Chrome(executable_path=PATH, options=option)

#---讀取單字表---#
wb = load_workbook('keywords.xlsx')
sheet = wb[data]
for k in range(1, sheet.max_row+1):
    if(str(sheet.cell(k, 1).value)!="None"):
        keywords.append(str(sheet.cell(k, 1).value))
allnum=str(len(keywords))
print("讀取到"+allnum+"個單字")

#增加超連結的函數
def add_hyperlink(paragraph, url, text, color, underline):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)
    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)
    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

#轉文字成XML相容
def valid_xml_char_ordinal(c):
    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
    )


#創建word
doc_initial = docx.Document()
# doc.add_heading('宥任自製題庫', level=1)
doc_initial.sections[0].header.paragraphs[0].text = "自動生成"+data+"題庫程式 by M110陳宥任。版權歸題目原作者所有。"
doc_initial.save(data+'.docx')
doc = docx.Document(data+'.docx')

#主程式 開始讀取
def QbankStart(QbankStartUrl):
    global allquesnum, checknum, Qrepeatcheck, repeatnum, longquesnum, quetag, valid_tags, invalidtagnum
    longquesnum = 0 #歸零
    repeatnum = 0 #歸零
    invalidtagnum = 0 #歸零
    driver.get(QbankStartUrl)
    #
    time.sleep(cooldown)
    #
    doc.add_heading(kw, level=2)
    ###
    Qtexts = driver.find_elements(By.CLASS_NAME, 'itemcontent')
    Anstexts = driver.find_elements(By.XPATH, '//*[@class="alert alert-success"]/div/div[1]/b[1]/a')
    Links = [Anstext.get_attribute('href') for Anstext in Anstexts]
    Qtags = driver.find_elements(By.XPATH, '//*[@class="col-lg-12 reponse-card"]/div[2]/div[@style="text-align:right"]')
    # for Qtext in Qtexts:
    #     print(Qtext.text)
    # for Anstext in Anstexts:
    #     print(Anstext.text)
    if(len(Qtexts)!=len(Anstexts)):
        print("題目與答案數目不一致!!!")
        driver.quit()
    allquesnum = len(Qtexts)
    for Qtext in Qtexts:
        tagbool=False #標籤判斷是否合法 預設無
        if Links[Qtexts.index(Qtext)] in Qrepeatcheck:
            # print("已重複")
            repeatnum+=1
            # doc.add_paragraph("已重複")
        elif len(str(Qtext.text)) > 1000:
            longquesnum+=1
            # add_hyperlink(doc.add_paragraph(), Links[Qtexts.index(Qtext)], "自動刪除長文", None, True)
            # doc.add_paragraph("-----")
        else:
            #-----#
            quetag = ''.join(c for c in str(Qtags[Qtexts.index(Qtext)].text) if valid_xml_char_ordinal(c))
            for tag in valid_tags:
                if tag in quetag: tagbool=True
            if tagbool:
                Qrepeatcheck.append(Links[Qtexts.index(Qtext)])
                doc.add_paragraph(''.join(c for c in str(Qtext.text) if valid_xml_char_ordinal(c)))
                add_hyperlink(doc.add_paragraph(), Links[Qtexts.index(Qtext)], "答案: "+Anstexts[Qtexts.index(Qtext)].text+"(解析連結)", None, True)
                doc.add_paragraph("-----")
                # print("符合標籤")
            else: 
                invalidtagnum +=1
                print("!!!已刪除-----> "+quetag)
            #-----#                    
    checknum+=1
    

for kw in keywords:
    if int(allnum) > 1 and checknum == int(allnum)//2:
        doc.save(data+'.docx')
        print("!!!--------進度 50% 存檔--------!!!")
        time.sleep(1.0)
        doc = docx.Document(data+'.docx')
    QbankStart("https://yamol.tw/tfulltext-"+str(kw)+".htm")
    time.sleep(0.3)
    print("完成"+kw+" 進度:"+str(checknum)+"/"+str(allnum)+" 一共"+str(allquesnum)+"題: "+str(repeatnum)+"重複,"+str(longquesnum)+"長文,"+str(invalidtagnum)+"非主題")

#全部標楷體化
# for paragraph in doc.paragraphs:
#     # if paragraph.style.name.startswith('Heading'):
#     for run in paragraph.runs:
#         # run.font.color.rgb = RGBColor(18, 255, 0)
#         run.font.name = '標楷體'
#         run._element.rPr.rFonts.set(qn('w:eastAsia'), u'標楷體')
#         # run.font.size = Pt(10)
#         # print(run.text)

doc.save(data+'.docx')
driver.quit()
print("完成"+str(checknum)+"個單字")
